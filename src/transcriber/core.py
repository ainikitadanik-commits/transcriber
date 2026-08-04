from __future__ import annotations

import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import wave
from array import array
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable


MODEL_NAME = "v3_e2e_rnnt"
SUPPORTED_MODELS = {"v3_e2e_rnnt", "v3_e2e_ctc"}
ENHANCEMENT_MODES = {"auto", "on", "off"}
DIARIZATION_MODEL = "pyannote/speaker-diarization-community-1"
SUPPORTED_EXTENSIONS = {
    ".aac",
    ".flac",
    ".m4a",
    ".mp3",
    ".mp4",
    ".ogg",
    ".wav",
    ".webm",
}


class TranscriptionError(RuntimeError):
    """An expected, user-facing transcription failure."""


@dataclass(frozen=True)
class Word:
    start: float
    end: float
    text: str


@dataclass(frozen=True)
class Segment:
    start: float
    end: float
    text: str
    speaker: str | None = None
    words: tuple[Word, ...] = ()


@dataclass(frozen=True)
class SpeakerTurn:
    start: float
    end: float
    label: str


def validate_input(input_path: Path) -> Path:
    path = input_path.expanduser().resolve()
    if not path.is_file():
        raise TranscriptionError(f"Файл не найден: {path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise TranscriptionError(
            f"Формат {path.suffix or '<без расширения>'} не поддерживается. "
            f"Допустимые форматы: {allowed}"
        )
    return path


def require_ffmpeg() -> str:
    executable = shutil.which("ffmpeg")
    if executable is None:
        raise TranscriptionError(
            "FFmpeg не найден. Установите FFmpeg и повторите команду."
        )
    return executable


def normalize_audio(
    input_path: Path, output_path: Path, ffmpeg: str, enhance_audio: bool = False
) -> None:
    command = [
        ffmpeg,
        "-hide_banner",
        "-loglevel",
        "error",
        "-y",
        "-i",
        str(input_path),
        "-vn",
    ]
    if enhance_audio:
        command.extend(
            ["-af", "highpass=f=70,loudnorm=I=-18:LRA=11:TP=-1.5"]
        )
    command.extend(
        [
            "-ac",
            "1",
            "-ar",
            "16000",
            "-c:a",
            "pcm_s16le",
            str(output_path),
        ]
    )
    try:
        subprocess.run(command, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as error:
        details = (error.stderr or error.stdout or "").strip()
        suffix = f" Детали FFmpeg: {details}" if details else ""
        raise TranscriptionError(f"Не удалось извлечь аудио.{suffix}") from error


def measure_loudness(input_path: Path, ffmpeg: str) -> float | None:
    command = [
        ffmpeg,
        "-hide_banner",
        "-nostats",
        "-i",
        str(input_path),
        "-af",
        "ebur128=framelog=quiet",
        "-f",
        "null",
        "-",
    ]
    result = subprocess.run(command, capture_output=True, text=True)
    matches = re.findall(r"I:\s*(-?\d+(?:\.\d+)?)\s+LUFS", result.stderr or "")
    return float(matches[-1]) if matches else None


def resolve_audio_enhancement(
    mode: str, input_path: Path, ffmpeg: str
) -> tuple[bool, float | None]:
    if mode not in ENHANCEMENT_MODES:
        raise TranscriptionError(f"Неизвестный режим тихих голосов: {mode}")
    if mode == "on":
        return True, None
    if mode == "off":
        return False, None
    loudness = measure_loudness(input_path, ffmpeg)
    return bool(loudness is not None and loudness < -20.0), loudness


def audio_duration(audio_path: Path) -> float:
    with wave.open(str(audio_path), "rb") as audio:
        return audio.getnframes() / audio.getframerate()


def combine_wav_parts(part_paths: list[Path], output_path: Path) -> None:
    if not part_paths:
        raise TranscriptionError("Не выбрано ни одной части встречи.")
    with wave.open(str(part_paths[0]), "rb") as first:
        params = first.getparams()
        format_key = (
            first.getnchannels(),
            first.getsampwidth(),
            first.getframerate(),
            first.getcomptype(),
        )
    with wave.open(str(output_path), "wb") as combined:
        combined.setparams(params)
        for part_path in part_paths:
            with wave.open(str(part_path), "rb") as part:
                part_format_key = (
                    part.getnchannels(),
                    part.getsampwidth(),
                    part.getframerate(),
                    part.getcomptype(),
                )
                if part_format_key != format_key:
                    raise TranscriptionError(
                        "Не удалось объединить части встречи: параметры аудио различаются."
                    )
                while content := part.readframes(1_000_000):
                    combined.writeframes(content)


def find_recovery_gaps(
    segments: list[Segment], duration: float, minimum_duration: float = 3.0
) -> list[tuple[float, float]]:
    gaps: list[tuple[float, float]] = []
    cursor = 0.0
    for segment in sorted(segments, key=lambda item: item.start):
        if segment.start - cursor >= minimum_duration:
            gaps.append((cursor, segment.start))
        cursor = max(cursor, segment.end)
    if duration - cursor >= minimum_duration:
        gaps.append((cursor, duration))
    return gaps


def gap_has_signal(
    audio_path: Path, start: float, end: float, minimum_dbfs: float = -52.0
) -> bool:
    with wave.open(str(audio_path), "rb") as audio:
        sample_width = audio.getsampwidth()
        if sample_width != 2:
            return True
        rate = audio.getframerate()
        audio.setpos(min(audio.getnframes(), max(0, int(start * rate))))
        raw = audio.readframes(max(0, int((end - start) * rate)))
    samples = array("h")
    samples.frombytes(raw)
    if sys.byteorder == "big":
        samples.byteswap()
    if not samples:
        return False
    rms = math.sqrt(sum(sample * sample for sample in samples) / len(samples))
    dbfs = 20 * math.log10(max(rms / 32768, 1e-9))
    return dbfs >= minimum_dbfs


def extract_audio_chunk(
    audio_path: Path,
    output_path: Path,
    start: float,
    end: float,
    ffmpeg: str,
) -> None:
    command = [
        ffmpeg,
        "-hide_banner",
        "-loglevel",
        "error",
        "-y",
        "-ss",
        f"{start:.3f}",
        "-i",
        str(audio_path),
        "-t",
        f"{end - start:.3f}",
        "-c:a",
        "pcm_s16le",
        str(output_path),
    ]
    subprocess.run(command, check=True, capture_output=True, text=True)


def extract_wav_chunk(
    audio_path: Path,
    output_path: Path,
    start: float,
    end: float,
) -> None:
    with wave.open(str(audio_path), "rb") as source:
        rate = source.getframerate()
        start_frame = min(source.getnframes(), max(0, int(start * rate)))
        frame_count = max(0, int((end - start) * rate))
        source.setpos(start_frame)
        content = source.readframes(frame_count)
        params = source.getparams()
    with wave.open(str(output_path), "wb") as target:
        target.setparams(params)
        target.writeframes(content)


def _remove_text_overlap(previous: str, current: str) -> str:
    previous_tokens = previous.split()
    current_tokens = current.split()
    maximum = min(len(previous_tokens), len(current_tokens), 16)
    for size in range(maximum, 0, -1):
        previous_tail = [token.casefold() for token in previous_tokens[-size:]]
        current_head = [token.casefold() for token in current_tokens[:size]]
        if previous_tail == current_head:
            return " ".join(current_tokens[size:])
    return current


def transcribe_audio_windows(
    audio_path: Path,
    device: str,
    maximum_window: float = 20.0,
    context: float = 0.75,
    model_loader: Callable[[str], Any] | None = None,
) -> list[Segment]:
    if maximum_window <= 0 or maximum_window + 2 * context > 25:
        raise ValueError("Окно распознавания должно помещаться в лимит 25 секунд.")

    model = (model_loader or load_model)(device)
    duration = audio_duration(audio_path)
    segments: list[Segment] = []
    core_start = 0.0
    window_index = 0
    while core_start < duration:
        core_end = min(duration, core_start + maximum_window)
        extract_start = max(0.0, core_start - context)
        extract_end = min(duration, core_end + context)
        chunk_path = audio_path.parent / f"window-{window_index:04}.wav"
        extract_wav_chunk(audio_path, chunk_path, extract_start, extract_end)
        result = model.transcribe(str(chunk_path), word_timestamps=True)
        words = tuple(
            Word(
                float(word.start) + extract_start,
                float(word.end) + extract_start,
                str(word.text).strip(),
            )
            for word in (getattr(result, "words", None) or [])
            if str(word.text).strip()
            and core_start - 0.02
            <= (float(word.start) + float(word.end)) / 2 + extract_start
            < core_end + 0.02
        )
        if words:
            text = _join_words(list(words))
            segments.append(
                Segment(words[0].start, words[-1].end, text, words=words)
            )
        else:
            text = str(getattr(result, "text", "")).strip()
            if text:
                if segments:
                    text = _remove_text_overlap(segments[-1].text, text)
                if text:
                    segments.append(Segment(core_start, core_end, text))
        core_start = core_end
        window_index += 1
    return segments


def recover_missing_segments(
    model: Any,
    audio_path: Path,
    existing_segments: list[Segment],
    ffmpeg: str,
    maximum_window: float = 20.0,
    context: float = 0.75,
) -> list[Segment]:
    duration = audio_duration(audio_path)
    gaps = find_recovery_gaps(existing_segments, duration)
    recovered: list[Segment] = []
    for gap_index, (gap_start, gap_end) in enumerate(gaps):
        if not gap_has_signal(audio_path, gap_start, gap_end):
            continue
        core_start = gap_start
        window_index = 0
        while core_start < gap_end:
            core_end = min(gap_end, core_start + maximum_window)
            extract_start = max(0.0, core_start - context)
            extract_end = min(duration, core_end + context)
            chunk_path = audio_path.parent / f"recovery-{gap_index}-{window_index}.wav"
            extract_audio_chunk(
                audio_path, chunk_path, extract_start, extract_end, ffmpeg
            )
            result = model.transcribe(str(chunk_path), word_timestamps=True)
            words = tuple(
                Word(
                    float(word.start) + extract_start,
                    float(word.end) + extract_start,
                    str(word.text).strip(),
                )
                for word in (result.words or [])
                if str(word.text).strip()
                and float(word.start) + extract_start >= core_start - 0.02
                and float(word.end) + extract_start <= core_end + 0.02
            )
            if words:
                recovered.append(
                    Segment(words[0].start, words[-1].end, _join_words(list(words)), words=words)
                )
            core_start = core_end
            window_index += 1
    return recovered


def choose_device(requested: str) -> str:
    if requested == "cpu":
        return "cpu"
    try:
        import torch
    except ImportError as error:
        raise TranscriptionError(
            "ML-зависимости не установлены. Выполните шаги установки из README."
        ) from error
    mps_available = bool(
        getattr(torch.backends, "mps", None) and torch.backends.mps.is_available()
    )
    if requested == "mps" and not mps_available:
        raise TranscriptionError("MPS недоступен в текущей сборке PyTorch.")
    return "mps" if mps_available else "cpu"


def pyannote_audio(audio_path: Path) -> dict[str, Any]:
    from gigaam.preprocess import SAMPLE_RATE, load_audio

    waveform = load_audio(str(audio_path)).unsqueeze(0)
    return {"waveform": waveform, "sample_rate": SAMPLE_RATE}


def _use_in_memory_audio_for_gigaam() -> None:
    from gigaam import vad_utils

    if getattr(vad_utils, "_transcriber_in_memory_audio", False):
        return
    original_get_pipeline = vad_utils.get_pipeline

    class InMemoryPipeline:
        def __init__(self, pipeline: Any):
            self.pipeline = pipeline

        def __call__(self, audio_path: str) -> Any:
            return self.pipeline(pyannote_audio(Path(audio_path)))

    def get_pipeline(device: Any, model_id: str = "pyannote/segmentation-3.0") -> Any:
        return InMemoryPipeline(original_get_pipeline(device, model_id))

    vad_utils.get_pipeline = get_pipeline
    vad_utils._transcriber_in_memory_audio = True


def load_model(device: str, model_name: str = MODEL_NAME) -> Any:
    try:
        import gigaam
    except ImportError as error:
        raise TranscriptionError(
            "GigaAM не установлен. Выполните шаги установки из README."
        ) from error
    _use_in_memory_audio_for_gigaam()
    model_dir = os.getenv("TRANSCRIBER_GIGAAM_MODELS_DIR")
    options: dict[str, Any] = {}
    if model_dir:
        root = Path(model_dir).expanduser()
        required = [root / f"{model_name}.ckpt"]
        if "e2e" in model_name:
            required.append(root / f"{model_name}_tokenizer.model")
        missing = [path.name for path in required if not path.is_file()]
        if missing:
            legacy_root = Path.home() / ".cache" / "gigaam"
            legacy_required = [legacy_root / path.name for path in required]
            if all(path.is_file() for path in legacy_required):
                root = legacy_root
                missing = []
        if missing:
            raise TranscriptionError(
                "Локальные веса GigaAM не найдены: " + ", ".join(missing)
            )
        options["download_root"] = str(root)
    return gigaam.load_model(
        model_name,
        device=device,
        fp16_encoder=device != "cpu",
        **options,
    )


def transcribe_audio(
    audio_path: Path,
    device: str,
    batch_size: int,
    word_timestamps: bool = False,
    model_loader: Callable[[str], Any] = load_model,
) -> list[Segment]:
    model = model_loader(device)
    try:
        result = model.transcribe_longform(
            str(audio_path),
            fr_batch_size=batch_size,
            word_timestamps=word_timestamps,
        )
    except RuntimeError as error:
        if "HF_TOKEN" in str(error) and "not found locally" in str(error):
            raise TranscriptionError(
                "VAD-модель pyannote не найдена в кэше. Примите условия "
                "pyannote/segmentation-3.0 и задайте HF_TOKEN только на время "
                "первой загрузки."
            ) from error
        raise
    return [
        Segment(
            float(item.start),
            float(item.end),
            str(item.text).strip(),
            words=tuple(
                Word(float(word.start), float(word.end), str(word.text).strip())
                for word in (getattr(item, "words", None) or [])
                if str(word.text).strip()
            ),
        )
        for item in result
        if str(item.text).strip()
    ]


def load_diarization_pipeline(device: str) -> Any:
    try:
        import torch
        from pyannote.audio import Pipeline

        from .models import DIARIZATION_MODEL, local_snapshot
    except ImportError as error:
        raise TranscriptionError(
            "Pyannote не установлен. Выполните шаги установки из README."
        ) from error
    try:
        pipeline = Pipeline.from_pretrained(local_snapshot(DIARIZATION_MODEL))
    except Exception as error:
        raise TranscriptionError(
            "Локальная модель разделения по спикерам не найдена. "
            "Введите Hugging Face Read-токен и повторите попытку."
        ) from error
    if pipeline is None:
        raise TranscriptionError(
            "Не удалось загрузить модель разделения по спикерам. Примите "
            "условия pyannote/speaker-diarization-community-1 и введите "
            "Hugging Face Read-токен для первого запуска."
        )
    return pipeline.to(torch.device(device))


def diarize_audio(
    audio_path: Path,
    device: str,
    num_speakers: int | None = None,
    pipeline_loader: Callable[[str], Any] = load_diarization_pipeline,
) -> list[SpeakerTurn]:
    pipeline = pipeline_loader(device)
    options = {"num_speakers": num_speakers} if num_speakers else {}
    output = pipeline(pyannote_audio(audio_path), **options)
    annotation = output.exclusive_speaker_diarization
    return [
        SpeakerTurn(float(turn.start), float(turn.end), str(label))
        for turn, _, label in annotation.itertracks(yield_label=True)
    ]


def _overlap(start: float, end: float, turn: SpeakerTurn) -> float:
    return max(0.0, min(end, turn.end) - max(start, turn.start))


def _speaker_for(start: float, end: float, turns: list[SpeakerTurn]) -> str | None:
    if not turns:
        return None
    best = max(turns, key=lambda turn: _overlap(start, end, turn))
    if _overlap(start, end, best) > 0:
        return best.label
    midpoint = (start + end) / 2
    return min(
        turns,
        key=lambda turn: min(abs(midpoint - turn.start), abs(midpoint - turn.end)),
    ).label


def _join_words(words: list[Word]) -> str:
    text = " ".join(word.text for word in words)
    text = re.sub(r"\s+([,.;:!?%…»)])", r"\1", text)
    return re.sub(r"([«(])\s+", r"\1", text).strip()


def assign_speakers(
    segments: list[Segment], turns: list[SpeakerTurn]
) -> list[Segment]:
    label_names: dict[str, str] = {}

    def display_name(label: str | None) -> str:
        if label is None:
            return "Спикер 1"
        if label not in label_names:
            label_names[label] = f"Спикер {len(label_names) + 1}"
        return label_names[label]

    words = [word for segment in segments for word in segment.words]
    if not words:
        return [
            Segment(
                segment.start,
                segment.end,
                segment.text,
                display_name(_speaker_for(segment.start, segment.end, turns)),
            )
            for segment in segments
        ]

    phrases: list[list[Word]] = []
    phrase: list[Word] = []
    for word in words:
        should_split = bool(
            phrase
            and (
                word.start - phrase[-1].end > 0.7
                or re.search(r"[.!?…][»\"']?$", phrase[-1].text)
                or phrase[-1].end - phrase[0].start >= 18.0
            )
        )
        if should_split:
            phrases.append(phrase)
            phrase = []
        phrase.append(word)
    if phrase:
        phrases.append(phrase)

    groups: list[tuple[str, list[Word]]] = []
    for phrase in phrases:
        votes: dict[str | None, float] = {}
        for word in phrase:
            label = _speaker_for(word.start, word.end, turns)
            votes[label] = votes.get(label, 0.0) + max(0.05, word.end - word.start)
        label = max(votes, key=votes.get)
        speaker = display_name(label)
        if (
            groups
            and groups[-1][0] == speaker
            and phrase[0].start - groups[-1][1][-1].end <= 1.0
        ):
            groups[-1][1].extend(phrase)
        else:
            groups.append((speaker, phrase))
    return [
        Segment(
            group[0].start,
            group[-1].end,
            _join_words(group),
            speaker,
        )
        for speaker, group in groups
    ]


def _format_timestamp(seconds: float) -> str:
    milliseconds = max(0, round(seconds * 1000))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02}.{millis:03}"


def build_txt(segments: list[Segment]) -> str:
    lines = [
        f"[{_format_timestamp(segment.start)} --> {_format_timestamp(segment.end)}] "
        f"{f'{segment.speaker}: ' if segment.speaker else ''}{segment.text}"
        for segment in segments
    ]
    return "\n\n".join(lines) + ("\n" if lines else "")


def write_docx(output_path: Path, document: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as error:
        raise TranscriptionError(
            "Для экспорта DOCX не установлен python-docx. Повторите установку из README."
        ) from error

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("ТРАНСКРИПЦИЯ ВСТРЕЧИ  •  GIGAAM")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string("7B8681")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Страница ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor.from_string("7B8681")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run("ЛОКАЛЬНАЯ ТРАНСКРИПЦИЯ")
    kicker_run.bold = True
    kicker_run.font.name = "Calibri"
    kicker_run.font.size = Pt(9)
    kicker_run.font.color.rgb = RGBColor.from_string("176B51")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Транскрипция встречи")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("16201D")

    input_info = document["input"]
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(input_info["name"])
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("67726E")

    processing = document["processing"]
    diarization = processing["diarization"]
    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(3)
    metadata_label = metadata.add_run("Обработка: ")
    metadata_label.bold = True
    metadata.add_run(
        f"{processing['model']} • {processing['device_used'].upper()} • "
        f"{len(input_info['parts'])} файл(а)"
    )
    speakers = doc.add_paragraph()
    speakers.paragraph_format.space_after = Pt(14)
    speakers_label = speakers.add_run("Спикеры: ")
    speakers_label.bold = True
    speakers.add_run(
        str(diarization["num_speakers"])
        if diarization["enabled"]
        else "разделение отключено"
    )

    if len(input_info["parts"]) > 1:
        doc.add_heading("Части встречи", level=1)
        for part in input_info["parts"]:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            label = paragraph.add_run(f"Часть {part['index']}: ")
            label.bold = True
            paragraph.add_run(
                f"[{_format_timestamp(part['start'])} – "
                f"{_format_timestamp(part['end'])}] {part['name']}"
            )

    doc.add_heading("Транскрипция", level=1)
    for segment in document["segments"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.25
        timestamp = paragraph.add_run(
            f"[{_format_timestamp(segment['start'])} – "
            f"{_format_timestamp(segment['end'])}]  "
        )
        timestamp.bold = True
        timestamp.font.name = "Calibri"
        timestamp.font.size = Pt(9)
        timestamp.font.color.rgb = RGBColor.from_string("176B51")
        if segment.get("speaker"):
            speaker = paragraph.add_run(f"{segment['speaker']}: ")
            speaker.bold = True
            speaker.font.name = "Calibri"
            speaker.font.size = Pt(11)
            speaker.font.color.rgb = RGBColor.from_string("16201D")
        text_run = paragraph.add_run(segment["text"])
        text_run.font.name = "Calibri"
        text_run.font.size = Pt(11)
        text_run.font.color.rgb = RGBColor.from_string("16201D")

    doc.core_properties.title = "Транскрипция встречи"
    doc.core_properties.subject = "Локальная транскрипция GigaAM"
    doc.core_properties.author = "Локальный транскрибатор"
    doc.save(output_path)


def build_document(
    input_path: Path | list[Path],
    segments: list[Segment],
    requested_device: str,
    used_device: str,
    batch_size: int,
    fallback_used: bool,
    started_at: datetime,
    elapsed_seconds: float,
    diarization_enabled: bool = False,
    diarization_device: str | None = None,
    speaker_count_requested: int | None = None,
    recovery_enabled: bool = False,
    recovered_segments: int = 0,
    recovery_device: str | None = None,
    enhancement_mode: str = "off",
    enhancement_applied: bool = False,
    model_name: str = MODEL_NAME,
    local_windowing: bool = False,
    part_metadata: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    input_paths = [input_path] if isinstance(input_path, Path) else input_path
    primary = input_paths[0]
    is_multipart = len(input_paths) > 1
    parts = part_metadata or [
        {
            "index": index,
            "path": str(path),
            "name": path.name,
            "format": path.suffix.lower().lstrip("."),
            "size_bytes": path.stat().st_size,
        }
        for index, path in enumerate(input_paths, start=1)
    ]
    speakers = sorted(
        {segment.speaker for segment in segments if segment.speaker},
        key=lambda value: int(value.rsplit(" ", 1)[-1]),
    )
    return {
        "schema_version": "1.3",
        "input": {
            "path": None if is_multipart else str(primary),
            "name": (
                f"{primary.stem} — {len(input_paths)} части"
                if is_multipart
                else primary.name
            ),
            "format": "multi_part" if is_multipart else primary.suffix.lower().lstrip("."),
            "size_bytes": sum(path.stat().st_size for path in input_paths),
            "parts": parts,
        },
        "processing": {
            "model": model_name,
            "mode": "local_windows" if local_windowing else "longform",
            "device_requested": requested_device,
            "device_used": used_device,
            "cpu_fallback_used": fallback_used,
            "batch_size": batch_size,
            "audio": {
                "sample_rate_hz": 16000,
                "channels": 1,
                "codec": "pcm_s16le",
                "enhancement_mode": enhancement_mode,
                "enhancement_applied": enhancement_applied,
            },
            "diarization": {
                "enabled": diarization_enabled,
                "model": DIARIZATION_MODEL if diarization_enabled else None,
                "device_used": diarization_device,
                "num_speakers_requested": speaker_count_requested,
                "num_speakers": len(speakers) if diarization_enabled else None,
            },
            "gap_recovery": {
                "enabled": recovery_enabled,
                "device_used": recovery_device,
                "recovered_segments": recovered_segments,
            },
            "started_at": started_at.isoformat(),
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "elapsed_seconds": round(elapsed_seconds, 3),
        },
        "text": " ".join(segment.text for segment in segments),
        "segments": [
            {
                "start": round(segment.start, 3),
                "end": round(segment.end, 3),
                "speaker": segment.speaker,
                "text": segment.text,
            }
            for segment in segments
        ],
    }


def write_outputs(
    output_dir: Path, stem: str, txt: str, document: dict[str, Any]
) -> tuple[Path, Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    txt_path = output_dir / f"{stem}.txt"
    json_path = output_dir / f"{stem}.json"
    docx_path = output_dir / f"{stem}.docx"
    txt_path.write_text(txt, encoding="utf-8")
    json_path.write_text(
        json.dumps(document, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    write_docx(docx_path, document)
    return txt_path, json_path, docx_path


def run(
    input_path: Path | list[Path],
    output_dir: Path,
    requested_device: str = "auto",
    batch_size: int = 2,
    diarization: bool = False,
    speaker_count: int | None = None,
    recover_gaps: bool = True,
    enhancement_mode: str = "auto",
    model_name: str = MODEL_NAME,
    local_windowing: bool = False,
    model_loader: Callable[[str], Any] = load_model,
    diarization_loader: Callable[[str], Any] = load_diarization_pipeline,
    progress_callback: Callable[[int, str, str], None] | None = None,
) -> tuple[Path, Path, Path, str, bool]:
    def report(progress: int, stage: str, message: str) -> None:
        if progress_callback:
            progress_callback(progress, stage, message)

    report(2, "preparing", "Проверяем файлы встречи…")
    requested_paths = [input_path] if isinstance(input_path, Path) else input_path
    if not requested_paths:
        raise TranscriptionError("Не выбрано ни одной части встречи.")
    sources = [validate_input(path) for path in requested_paths]
    if model_name not in SUPPORTED_MODELS:
        raise TranscriptionError(f"Неизвестная модель GigaAM: {model_name}")
    if enhancement_mode not in ENHANCEMENT_MODES:
        raise TranscriptionError(
            f"Неизвестный режим тихих голосов: {enhancement_mode}"
        )
    ffmpeg = require_ffmpeg()
    device = choose_device(requested_device)
    started_at = datetime.now(timezone.utc)
    started = time.monotonic()
    fallback_used = False
    diarization_device: str | None = None
    recovery_device: str | None = None
    recovered_count = 0
    loaded_model: Any = None
    part_metadata: list[dict[str, Any]] = []
    enhancement_applied = False

    def tracked_model_loader(selected_device: str) -> Any:
        nonlocal loaded_model
        report(20, "loading_model", "Загружаем модель распознавания…")
        loaded_model = (
            load_model(selected_device, model_name)
            if model_loader is load_model
            else model_loader(selected_device)
        )
        report(28, "transcribing", "Распознаём речь и расставляем таймкоды…")
        return loaded_model

    with tempfile.TemporaryDirectory(prefix="transcriber-") as temp_dir:
        temp_path = Path(temp_dir)
        normalized_parts: list[Path] = []
        offset = 0.0
        for index, source in enumerate(sources, start=1):
            report(
                3 + round(index / len(sources) * 12),
                "preparing",
                f"Подготавливаем часть {index} из {len(sources)}…",
            )
            should_enhance, loudness = resolve_audio_enhancement(
                enhancement_mode, source, ffmpeg
            )
            normalized_part = temp_path / f"part-{index:03}.wav"
            normalize_audio(source, normalized_part, ffmpeg, should_enhance)
            duration = audio_duration(normalized_part)
            normalized_parts.append(normalized_part)
            enhancement_applied = enhancement_applied or should_enhance
            part_metadata.append(
                {
                    "index": index,
                    "path": str(source),
                    "name": source.name,
                    "format": source.suffix.lower().lstrip("."),
                    "size_bytes": source.stat().st_size,
                    "start": round(offset, 3),
                    "end": round(offset + duration, 3),
                    "duration_seconds": round(duration, 3),
                    "input_loudness_lufs": loudness,
                    "enhancement_applied": should_enhance,
                }
            )
            offset += duration

        if len(normalized_parts) == 1:
            normalized_audio = normalized_parts[0]
        else:
            normalized_audio = temp_path / "meeting.wav"
            combine_wav_parts(normalized_parts, normalized_audio)
        report(18, "preparing", "Аудио подготовлено, запускаем модель…")
        try:
            if local_windowing:
                segments = transcribe_audio_windows(
                    normalized_audio,
                    device,
                    model_loader=tracked_model_loader,
                )
            else:
                segments = transcribe_audio(
                    normalized_audio,
                    device,
                    batch_size,
                    word_timestamps=diarization,
                    model_loader=tracked_model_loader,
                )
        except TranscriptionError:
            raise
        except RuntimeError as error:
            if device != "mps":
                raise
            fallback_used = True
            device = "cpu"
            if local_windowing:
                segments = transcribe_audio_windows(
                    normalized_audio,
                    device,
                    model_loader=tracked_model_loader,
                )
            else:
                segments = transcribe_audio(
                    normalized_audio,
                    device,
                    batch_size,
                    word_timestamps=diarization,
                    model_loader=tracked_model_loader,
                )

        report(72, "transcribing", "Основное распознавание завершено…")

        if recover_gaps:
            report(74, "recovery", "Проверяем пропущенные интервалы…")
            recovery_device = device
            try:
                recovered = recover_missing_segments(
                    loaded_model, normalized_audio, segments, ffmpeg
                )
            except (RuntimeError, NotImplementedError):
                if recovery_device != "mps":
                    raise
                fallback_used = True
                recovery_device = "cpu"
                cpu_model = (
                    load_model("cpu", model_name)
                    if model_loader is load_model
                    else model_loader("cpu")
                )
                recovered = recover_missing_segments(
                    cpu_model, normalized_audio, segments, ffmpeg
                )
            recovered_count = len(recovered)
            segments = sorted([*segments, *recovered], key=lambda item: item.start)
            report(84, "recovery", "Проверка пропусков завершена…")

        if diarization:
            report(86, "diarization", "Разделяем реплики по спикерам…")
            diarization_device = device
            try:
                turns = diarize_audio(
                    normalized_audio,
                    diarization_device,
                    speaker_count,
                    pipeline_loader=diarization_loader,
                )
            except (RuntimeError, NotImplementedError):
                if diarization_device != "mps":
                    raise
                fallback_used = True
                diarization_device = "cpu"
                turns = diarize_audio(
                    normalized_audio,
                    "cpu",
                    speaker_count,
                    pipeline_loader=diarization_loader,
                )
            segments = assign_speakers(segments, turns)
            report(94, "diarization", "Разделение по спикерам завершено…")

    if not segments:
        raise TranscriptionError(
            "GigaAM не обнаружил речь в файле; результаты не созданы."
        )
    report(97, "exporting", "Формируем TXT, JSON и DOCX…")
    elapsed = time.monotonic() - started
    document = build_document(
        input_path=sources,
        segments=segments,
        requested_device=requested_device,
        used_device=device,
        batch_size=batch_size,
        fallback_used=fallback_used,
        started_at=started_at,
        elapsed_seconds=elapsed,
        diarization_enabled=diarization,
        diarization_device=diarization_device,
        speaker_count_requested=speaker_count,
        recovery_enabled=recover_gaps,
        recovered_segments=recovered_count,
        recovery_device=recovery_device,
        enhancement_mode=enhancement_mode,
        enhancement_applied=enhancement_applied,
        model_name=model_name,
        local_windowing=local_windowing,
        part_metadata=part_metadata,
    )
    output_stem = sources[0].stem
    if len(sources) > 1:
        output_stem += "-combined"
    if diarization:
        output_stem += "-speakers"
    txt_path, json_path, docx_path = write_outputs(
        output_dir.expanduser().resolve(), output_stem, build_txt(segments), document
    )
    return txt_path, json_path, docx_path, device, fallback_used
